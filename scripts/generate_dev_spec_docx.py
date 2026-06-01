#!/usr/bin/env python3
"""Generate development specification as DOCX from project analysis."""
from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DEFAULT_REPO = Path(".")
DEFAULT_OUTPUT = Path("docs/AutoBBoxToolkit_DevSpec.docx")


def load_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def load_jsonl(path: Path) -> list[dict]:
    if not path.exists():
        return []
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def extract_feature_list(repo: Path) -> list[dict]:
    """Extract feature list from feature_index.json and msg text."""
    features = []
    fi = load_json(repo / ".autobbox/index/feature_index.json", {"features": []})
    msg_path = repo / "autobbox_msg.txt"

    # Parse message text for command definitions
    msg_entries: dict[str, dict] = {}
    if msg_path.exists():
        lines = msg_path.read_text(encoding="utf-8", errors="replace").splitlines()
        current_key = None
        for line in lines:
            line = line.strip()
            if not line or line == "#":
                continue
            # First line after separator is the key
            if re.match(r"^[A-Za-z][A-Za-z0-9_]+$", line) and len(line) > 3:
                current_key = line
                msg_entries[current_key] = {"key": current_key, "label": "", "tip": ""}
            elif current_key:
                if not msg_entries[current_key]["label"]:
                    msg_entries[current_key]["label"] = line
                elif not msg_entries[current_key]["tip"]:
                    msg_entries[current_key]["tip"] = line

    for feat in fi.get("features", []):
        fid = feat.get("feature_id", "unknown")
        title = feat.get("title", fid)
        paths = feat.get("paths", {})
        commands = feat.get("commands", [])

        features.append({
            "id": fid,
            "title": title,
            "commands": commands,
            "paths": {k: v if isinstance(v, list) else [v] for k, v in (paths or {}).items()},
            "notes": feat.get("notes", ""),
        })

    return features


def generate_docx(repo: Path, output: Path) -> None:
    doc = Document()

    # --- Styles ---
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)

    # === Cover Page ===
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AutoBBoxToolkit")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Creo TOOLKIT Plugin Development Specification")
    run.font.size = Pt(16)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Version: 1.0\nDate: 2026-05-28\nTarget: PTC Creo 10.0.8.0").font.size = Pt(10)

    doc.add_page_break()

    # === TOC placeholder ===
    doc.add_heading("Table of Contents", level=0)
    doc.add_paragraph("(Generated automatically by Word — press Ctrl+A then F9 to update)")

    doc.add_page_break()

    # === 1. Overview ===
    doc.add_heading("1. Project Overview", level=0)
    doc.add_paragraph(
        "AutoBBoxToolkit is a Creo TOOLKIT DLL plugin that extends PTC Creo Parametric "
        "with automated CAD workflows. It provides 18+ registered commands covering "
        "dimension annotation, BOM management, drawing view creation, family table "
        "management, simplified representation creation, and assembly utilities."
    )
    doc.add_paragraph(
        "The plugin is built as a C++17 shared library using CMake, targeting Creo 10.0.8.0 "
        "on Windows x64 with MSVC. All source code follows a layered architecture: "
        "main (orchestration), application (business logic), ui (dialogs), creo (API wrappers), "
        "and common (utilities)."
    )

    # === 2. Architecture ===
    doc.add_heading("2. Architecture", level=0)

    doc.add_heading("2.1 Module Layout", level=1)
    arch_table = doc.add_table(rows=7, cols=3)
    arch_table.style = "Light Grid Accent 1"
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (layer, dir_path, desc) in enumerate([
        ("Entry", "src/autobbox_toolkit.cpp", "Thin plugin shell (~252 lines), user_initialize/terminate"),
        ("Main", "src/main/", "Command registration, callback wiring, runtime bridge"),
        ("Application", "src/application/", "Feature workflows and business logic"),
        ("UI", "src/ui/", "Native Creo dialog/controller logic"),
        ("Creo", "src/creo/", "Creo TOOLKIT API wrapper helpers"),
        ("Common", "src/common/", "Cross-cutting utilities: log, string, file"),
        ("Headers", "include/autobbox/", "114 header files, 1:1 with source modules"),
    ]):
        arch_table.rows[i].cells[0].text = layer
        arch_table.rows[i].cells[1].text = dir_path
        arch_table.rows[i].cells[2].text = desc

    doc.add_paragraph()
    doc.add_heading("2.2 Build System", level=1)
    doc.add_paragraph(
        "CMake 3.20+ with MSVC (Visual Studio 2022). Links against Creo TOOLKIT libraries: "
        "protk_dllmd_NU.lib, ucore.lib, udata.lib, ws2_32.lib, wsock32.lib, mpr.lib, netapi32.lib. "
        "Output DLL: deploy/AutoBBoxToolkit/autobbox_toolkit.dll"
    )

    # === 3. Features ===
    doc.add_heading("3. Feature Catalog", level=0)
    features = extract_feature_list(repo)

    feat_table = doc.add_table(rows=1, cols=4)
    feat_table.style = "Light Grid Accent 1"
    hdr = feat_table.rows[0].cells
    hdr[0].text = "Feature ID"
    hdr[1].text = "Title"
    hdr[2].text = "Commands"
    hdr[3].text = "Status"

    for feat in features[:30]:
        row = feat_table.add_row()
        row.cells[0].text = feat["id"]
        row.cells[1].text = feat["title"]
        row.cells[2].text = ", ".join(feat["commands"][:3])
        row.cells[3].text = "Implemented"

    doc.add_page_break()

    # === 4. API Index ===
    doc.add_heading("4. Key API Reference", level=0)
    doc.add_paragraph("Relevant Creo TOOLKIT APIs organized by category:")

    api_index = load_json(repo / "docs/api_extracts/api_quickref.json", {"apis": {}})
    apis = api_index.get("apis", {})

    # Group by category
    by_cat: dict[str, list[str]] = defaultdict(list)
    for api_name, info in apis.items():
        by_cat[info.get("category", "Other")].append(api_name)

    for cat, api_list in sorted(by_cat.items()):
        if cat == "Other":
            continue
        doc.add_heading(f"4.{list(by_cat).index(cat) + 1} {cat}", level=1)
        cat_table = doc.add_table(rows=1, cols=2)
        cat_table.style = "Light Grid Accent 1"
        cat_table.rows[0].cells[0].text = "API"
        cat_table.rows[0].cells[1].text = "Synopsis"
        for api_name in api_list[:15]:
            info = apis.get(api_name, {})
            row = cat_table.add_row()
            row.cells[0].text = api_name
            row.cells[1].text = info.get("synopsis", "")[:120]

    doc.add_page_break()

    # === 5. Resource Layout ===
    doc.add_heading("5. Resource & Deployment", level=0)
    doc.add_paragraph(
        "Resources are organized under:\n"
        "  - resource/  — .res dialog definitions (25 files)\n"
        "  - ribbon/    — toolkitribbonui.rbn\n"
        "  - text/      — autobbox_msg.txt (message keys + Chinese labels)\n"
        "  - deploy/AutoBBoxToolkit/  — runtime output (DLL + mirrored resources)\n"
        "  - runtime/AutoBBoxToolkit/ — runtime mirror"
    )

    # === 6. Scripts ===
    doc.add_heading("6. Development Scripts", level=0)
    scripts = [
        ("build_autobbox.ps1", "Full build: cmake configure + build + resource sync"),
        ("backup_autobbox.ps1", "Backup source + runtime plugin as zip"),
        ("find_creo_context.ps1", "4-layer Creo API context lookup"),
        ("update_creo_install_index.ps1", "Rebuild Creo install search index"),
        ("update_creo_evidence_cache.ps1", "Rebuild project evidence cache + API graph"),
        ("build_api_graph.py", "Build API co-occurrence graph from project source"),
        ("generate_api_knowledge.py", "Extract API docs from protkdoc HTML → .h + .md"),
        ("merge_learned_aliases.py", "Merge learned query aliases into alias map"),
    ]
    script_table = doc.add_table(rows=1, cols=2)
    script_table.style = "Light Grid Accent 1"
    script_table.rows[0].cells[0].text = "Script"
    script_table.rows[0].cells[1].text = "Description"
    for name, desc in scripts:
        row = script_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = desc

    doc.add_paragraph()
    doc.add_paragraph(
        "Document auto-generated by scripts/generate_dev_spec_docx.py\n"
        "Source: feature_index.json, api_quickref.json, autobbox_msg.txt"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"DOCX written: {output}")


def main() -> int:
    repo = DEFAULT_REPO.resolve()
    output = DEFAULT_OUTPUT
    generate_docx(repo, output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
